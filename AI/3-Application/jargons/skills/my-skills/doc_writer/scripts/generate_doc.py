#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_doc.py

doc_writer skill 的核心脚本，将 JSON 结构化内容转换为 .docx 文档。

用法：
    python generate_doc.py --input content.json --output output.docx
    python generate_doc.py --input content.json            # 默认输出到 ../outputs/<title>.docx
    echo '{"title":"..."}' | python generate_doc.py --stdin

依赖：
    pip install python-docx
"""

import argparse
import json
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print(
        "错误：缺少 python-docx 库，请先运行：pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


# ─── 字体 & 尺寸常量 ──────────────────────────────────────────────────────────

FONT_CN_BODY  = "宋体"
FONT_EN_BODY  = "Times New Roman"
FONT_CN_TITLE = "黑体"
FONT_EN_TITLE = "Arial"

PT_H1   = 16   # 三号 ≈ 16pt
PT_H2   = 14   # 四号 ≈ 14pt
PT_H3   = 12   # 小四 ≈ 12pt
PT_BODY = 12   # 正文 12pt


# ─── 页面设置 ─────────────────────────────────────────────────────────────────

def _set_page_layout(doc: Document) -> None:
    """A4 页面，按 skill.md 规范设置页边距。"""
    sec = doc.sections[0]
    sec.page_height    = Cm(29.7)
    sec.page_width     = Cm(21.0)
    sec.top_margin     = Cm(2.54)
    sec.bottom_margin  = Cm(2.54)
    sec.left_margin    = Cm(3.18)
    sec.right_margin   = Cm(3.18)


# ─── 底层字体/间距工具 ────────────────────────────────────────────────────────

def _set_run_font(
    run,
    cn_font: str,
    en_font: str,
    size_pt: int,
    bold: bool = False,
) -> None:
    """同时设置英文字体和中文字体（东亚字体）。"""
    run.font.name = en_font
    run.font.bold = bold
    run.font.size = Pt(size_pt)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), cn_font)
    r_pr.insert(0, r_fonts)


def _set_para_spacing(
    para,
    before_pt: float = 6.0,
    after_pt: float = 6.0,
    line_multiple: float = None,
) -> None:
    """设置段前段后间距，以及多倍行距（line_multiple = 1.5 → 1.5 倍行距）。"""
    p_pr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"),  str(int(after_pt  * 20)))
    if line_multiple is not None:
        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:line"),     str(int(line_multiple * 240)))
    p_pr.append(spacing)


# ─── 内容构建工具 ─────────────────────────────────────────────────────────────

def add_heading(doc: Document, text: str, level: int):
    """写入 1~3 级标题段落。"""
    para = doc.add_heading(text, level=level)
    size_map = {1: PT_H1, 2: PT_H2, 3: PT_H3}
    size = size_map.get(level, PT_BODY)
    for run in para.runs:
        _set_run_font(run, FONT_CN_TITLE, FONT_EN_TITLE, size, bold=True)
    _set_para_spacing(
        para,
        before_pt=12 if level == 1 else 8,
        after_pt=6,
    )
    return para


def add_paragraph(doc: Document, text: str, first_line_indent_cm: float = 0.74):
    """写入正文段落（默认两字符首行缩进）。"""
    para = doc.add_paragraph()
    run  = para.add_run(text)
    _set_run_font(run, FONT_CN_BODY, FONT_EN_BODY, PT_BODY)
    _set_para_spacing(para, before_pt=3, after_pt=3, line_multiple=1.5)
    if first_line_indent_cm:
        para.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    return para


def add_bullet(doc: Document, text: str):
    """写入无序列表项。"""
    try:
        para = doc.add_paragraph(style="List Bullet")
    except KeyError:
        para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, FONT_CN_BODY, FONT_EN_BODY, PT_BODY)
    return para


def add_numbered(doc: Document, text: str):
    """写入有序列表项。"""
    try:
        para = doc.add_paragraph(style="List Number")
    except KeyError:
        para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, FONT_CN_BODY, FONT_EN_BODY, PT_BODY)
    return para


def add_table(doc: Document, headers: list, rows: list):
    """写入表格，第一行为加粗表头。"""
    col_count = len(headers)
    if col_count == 0:
        return
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # 表头行
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for run in hdr_cells[i].paragraphs[0].runs:
            _set_run_font(run, FONT_CN_BODY, FONT_EN_BODY, PT_BODY, bold=True)

    # 数据行
    for r_i, row in enumerate(rows):
        data_cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            if c_i < col_count:
                data_cells[c_i].text = str(val)

    return table


def add_page_number(doc: Document) -> None:
    """在页脚中间插入 PAGE 域（自动页码）。"""
    footer = doc.sections[0].footer
    para   = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run()

    for fld_type in ("begin", None, "end"):
        if fld_type in ("begin", "end"):
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), fld_type)
            run._r.append(fld)
        else:
            instr = OxmlElement("w:instrText")
            instr.text = "PAGE"
            run._r.append(instr)


def add_cover(
    doc: Document,
    title: str,
    subtitle: str = "",
    author: str = "",
    date: str = "",
    organization: str = "",
) -> None:
    """写入封面页并追加分页符。"""
    for _ in range(4):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p_title.add_run(title), FONT_CN_TITLE, FONT_EN_TITLE, 22, bold=True)

    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(p_sub.add_run(subtitle), FONT_CN_TITLE, FONT_EN_TITLE, 16)

    doc.add_paragraph()
    for label, value in [("作者", author), ("日期", date), ("单位", organization)]:
        if value:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font(
                p.add_run(f"{label}：{value}"),
                FONT_CN_BODY, FONT_EN_BODY, PT_BODY,
            )

    doc.add_page_break()


# ─── 简版 Markdown 解析 ───────────────────────────────────────────────────────

def _parse_markdown(doc: Document, md: str) -> None:
    """
    将 Markdown 文本逐行解析后写入 doc。
    支持：# 标题、- / * 无序列表、1. 有序列表、| 表格、普通段落。
    """
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line     = lines[i]
        stripped = line.strip()

        # 标题
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], level=3)
        elif stripped.startswith("## "):
            add_heading(doc, stripped[3:], level=2)
        elif stripped.startswith("# "):
            add_heading(doc, stripped[2:], level=1)

        # 无序列表
        elif stripped.startswith(("- ", "* ")):
            add_bullet(doc, stripped[2:])

        # 有序列表（形如 "1. "）
        elif (
            len(stripped) > 2
            and stripped[0].isdigit()
            and stripped[1] == "."
            and stripped[2:3] == " "
        ):
            add_numbered(doc, stripped[3:])

        # 表格（收集连续 | 开头的行）
        elif stripped.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            # 跳过 |---|---| 分隔行
            real_rows = [
                r for r in tbl_lines
                if not all(c in "| -:" for c in r)
            ]
            if real_rows:
                cells = [
                    [c.strip() for c in r.strip("|").split("|")]
                    for r in real_rows
                ]
                add_table(doc, cells[0], cells[1:])
            continue  # i 已在内层循环推进

        # 空行跳过
        elif not stripped:
            pass

        # 普通段落
        else:
            add_paragraph(doc, stripped)

        i += 1


# ─── 章节递归渲染 ─────────────────────────────────────────────────────────────

def _render_section(doc: Document, sec: dict, default_level: int = 2) -> None:
    """
    递归渲染单个 section 节点。
    支持字段：heading, level, content, markdown, bullets, numbered, table, subsections。
    """
    level = sec.get("level", default_level)

    if "heading" in sec:
        add_heading(doc, sec["heading"], level=level)

    if "content" in sec:
        add_paragraph(doc, sec["content"])

    if "markdown" in sec:
        _parse_markdown(doc, sec["markdown"])

    for item in sec.get("bullets", []):
        add_bullet(doc, item)

    for item in sec.get("numbered", []):
        add_numbered(doc, item)

    if "table" in sec:
        t = sec["table"]
        add_table(doc, t.get("headers", []), t.get("rows", []))

    for sub in sec.get("subsections", []):
        sub.setdefault("level", level + 1)
        _render_section(doc, sub, default_level=level + 1)


# ─── 主构建函数 ───────────────────────────────────────────────────────────────

def build_doc(data: dict) -> Document:
    """
    根据 JSON 数据构建并返回 Document 对象。

    JSON schema（所有字段均可选）：
    {
        "title":           str,          # 文档标题
        "document_type":   str,          # 类型标注（仅用于日志）
        "cover": {                        # 封面（有则加封面页）
            "subtitle":    str,
            "author":      str,
            "date":        str,
            "organization": str
        },
        "header_text":     str,          # 页眉文字
        "page_number":     bool,         # 是否添加页码，默认 true
        "abstract":        str,          # 摘要段落
        "sections": [ ... ]              # 章节列表，见 _render_section
    }
    """
    doc   = Document()
    title = data.get("title", "文档")

    _set_page_layout(doc)

    # 封面 or 顶部标题
    if "cover" in data:
        cv = data["cover"]
        add_cover(
            doc,
            title=title,
            subtitle=cv.get("subtitle", ""),
            author=cv.get("author", ""),
            date=cv.get("date", ""),
            organization=cv.get("organization", ""),
        )
    else:
        add_heading(doc, title, level=1)

    # 页眉
    if "header_text" in data:
        hpara = doc.sections[0].header.paragraphs[0]
        hpara.text = str(data["header_text"])
        hpara.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 页码
    if data.get("page_number", True):
        add_page_number(doc)

    # 摘要
    if "abstract" in data:
        add_heading(doc, "摘要", level=2)
        add_paragraph(doc, data["abstract"])

    # 正文章节
    for sec in data.get("sections", []):
        _render_section(doc, sec)

    return doc


# ─── CLI ─────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="将 JSON 结构内容生成为 .docx 文档（doc_writer skill 脚本）"
    )
    parser.add_argument(
        "--input", "-i",
        help="输入 JSON 文件路径（与 --stdin 二选一）",
    )
    parser.add_argument(
        "--stdin",
        action="store_true",
        help="从 stdin 读取 JSON 内容",
    )
    parser.add_argument(
        "--output", "-o",
        help="输出 .docx 文件路径（默认：../outputs/<title>.docx）",
    )
    args = parser.parse_args()

    # 读取输入
    if args.stdin:
        raw = sys.stdin.read()
    elif args.input:
        input_path = Path(args.input)
        if not input_path.exists():
            print(f"错误：输入文件不存在：{args.input}", file=sys.stderr)
            sys.exit(1)
        raw = input_path.read_text(encoding="utf-8")
    else:
        parser.print_help()
        sys.exit(1)

    try:
        data = json.loads(raw)
    except json.JSONDecodeError as exc:
        print(f"错误：JSON 解析失败 — {exc}", file=sys.stderr)
        sys.exit(1)

    # 构建文档
    doc = build_doc(data)

    # 确定输出路径
    if args.output:
        out_path = Path(args.output)
    else:
        safe_title = (
            data.get("title", "document")
            .replace("/", "-")
            .replace("\\", "-")
            .replace(" ", "_")
        )
        out_dir  = Path(__file__).resolve().parent.parent / "outputs"
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"{safe_title}.docx"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    result = {
        "status":        "success",
        "title":         data.get("title", ""),
        "document_type": data.get("document_type", ""),
        "file_name":     out_path.name,
        "output_path":   str(out_path.resolve()),
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
